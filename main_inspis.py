'''
generiranje zadataka - ukupan plan
1. random selekcija iz dokumenata [OK]
2. generiraj slucajne brojeve [OK]
3. rijesi formule [OK]
4. ispisi zadatke u docx [start, treba na kompu]
5. ispisi rijesenja u docx [start, treba na kompu]
6. program za unos zadataka [CURRENT]
7. graficko sucelje za programe
8. web sucelje za programe
9. accounti
10. mobilno sucelje za kreiranje testa i citanje rezultata
11. mobilno sucelje za ocjenjivanje?
'''

# redo as objects
from random import randint, random, choice
from math import log10
import sympy

def randData(type, frm, to, exp):
  if type.upper() == "N":
    x = randint(frm, to) * 10**exp
    if exp < 0 and x.is_integer():
      x = int(x)
    return round(x, abs(exp)+5)
  else:
    return (frm + (to - frm) * random()) * 10**exp 

def roundn(x, n):
  if x < 1:
    return abs(round(x, abs(int(log10(abs(x)))) + n))
  else:
    y = abs(int(log10(abs(x))))
    z = round(x / 10 ** y, n) * 10 ** y
    if float.is_integer(float(z)):
      return abs(int(z))
    else:
      return abs(z)

class Question:
  def __init__(self, data_dict, formulae):
    self.question = data_dict['question']
    self.constants = {i: constants[i] for i in data_dict['constants']}
    self.parameters = data_dict['parameters']
    self.unknown = data_dict['unknown']
    self.formulae = formulae + data_dict['formulae']
    self.knowns = data_dict['knowns']
    self.data = {}
    self.equations = [sympy.Eq(*map(sympy.S, formula.split("="))) for formula in self.formulae]
    self.symbols = {b.name: b for a in [eq.free_symbols for eq in self.equations] for b in a}
    self.generateData()
    self.solve()
  
  def generateData(self):
    for key, value in self.knowns.items():
      self.data[self.symbols[key]] = randData(*eval(value))
    if self.parameters:
      self.data.update({self.symbols[i]: self.parameters[i] for i in self.parameters})
    if self.constants:
      self.data.update({self.symbols[i]: self.constants[i] for i in self.constants})
    return
  
  def solve(self):
    while self.equations:
      for eq in self.equations:
        unknown = eq.free_symbols - set(self.data.keys())
        if len(unknown) == 1:
          solved_exp = sympy.solve(eq, list(unknown)[0])
          solved_data = [i.subs(self.data) for i in solved_exp]
          solved_data.sort(reverse = True)
          self.data[list(unknown)[0]] = solved_data[0]
          self.equations.remove(eq)
    res = self.data[self.symbols[self.unknown]]
    self.solution_exact = res
    self.solution_decimal = roundn(self.solution_exact, 3)
    return
 
  def __str__(self):
    return self.question.format(**{i:self.data[self.symbols[i]] for i in self.symbols})

with open('constants.txt', 'r') as consts:
  constants = {c.split("=")[0]: eval(c.split("=")[1]) for c in consts.readlines()}

#n = int(input("Koliko testova treba složiti? "))
n = 5
import docx
doc = docx.Document()
doc_res = docx.Document()

for test_number in range(n):
    qs = []
    for i in range(1,6):
      with open('z{}.txt'.format(i), 'r', encoding="utf-8") as zad:
        formulae, questions = zad.read().split('\n#questions\n')
        formulae = formulae.split('\n')[1:]
        questions = eval(choice(questions.split('\n')))
        q = Question(questions, formulae)
        print('{}. {}'.format(i, q))
        qs.append(q)
    print("\nSOLUTIONS")
    for i, q in enumerate(qs):
      print("{}. {}={}".format(i+1, q.unknown, q.solution_decimal))

    doc.add_heading('Zadaci', 1)
    doc.add_paragraph("Test#: {}".format(test_number + 1))
    doc_res.add_heading("Test#: {}".format(test_number + 1), 1)
    for i, q in enumerate(qs):
      doc.add_paragraph("{}. {}".format(i + 1, q))
      doc_res.add_paragraph("{}. {}={}".format(i+1, q.unknown, q.solution_decimal))
    doc.add_page_break()
doc.save('tests.docx')
doc_res.save('results.docx')


