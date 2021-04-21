# izlistati kategorije
# izlistati subkategorije
# izlistati zadatke
# odabrati zadatke

from _folders import *
from _problems import *

print("Odaberi podkategoriju (slovo) za test:")
selection = {}
i = ord('a')
for key, value in categories().items():
  print("{}:".format(key))
  for subcat in value:
    print("\t{}: {}".format(chr(i), subcat))
    selection[i] = [key, subcat]
    i += 1
letter = input("Odaberi slovo: ")
while ord(letter) not in list(selection.keys()):
  letter = input("Krivo slovo, odaberi opet: ")
path = "/".join(selection[ord(letter)])

problems = [i.name for i in os.scandir(path)]
print("Odaberi probleme koje želiš dodati u test:")
for index, problem in enumerate(problems):
  print(index+1, problem)
test = input("Unesi probleme odvojene razmakom: ").split()
while not all([i.isnumeric() for i in test]) or int(max(test)) > index + 1:
  test = input("Krivi unos, probaj ponovo: ")
test = [int(i) - 1 for i in test]
problems = [i.name for i in os.scandir(path)]
problems = [path + "/" + problems[i] for i in test]


#n = int(input("Koliko testova treba složiti? "))
n = 1
import docx
from random import *
doc = docx.Document()
doc_res = docx.Document()
for test_number in range(n):
    qs = []
    for i, problem in enumerate(problems):
      with open(problem, 'r', encoding="utf-8") as zad:
        formulae, questions = zad.read().split('\n#questions\n')
        formulae = formulae.split('\n')[1:]
        questions = eval(choice(questions.split('\n')))
        q = Question(questions, formulae)
        print('{}. {}'.format(i+1, q))
        qs.append(q)
    print("\nSOLUTIONS")
    for i, q in enumerate(qs):
      print("{}. {}={}".format(i+1, q.unknown, q.solution_decimal))

    doc.add_heading('Zadaci', 1)
    doc.add_paragraph("Test#: {}".format(test_number + 1))
    doc_res.add_heading("Test#: {}".format(test_number + 1), 1)
    for i, q in enumerate(qs):
      doc.add_paragraph("{}. {}".format(i + 1, q))
      doc_res.add_paragraph("{}. {}={}".format(i+1, q.unknown, q.solution_decimal))
    doc.add_page_break()
doc.save('test.docx')
doc_res.save('result.docx')
